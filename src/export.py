from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from docx import Document
from docx.shared import Inches
import io

def export_pdf(items, atn=None, citation_format="apa"):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    if atn:
        story.append(Paragraph(f"Assessment Task: {atn}", styles['Title']))
        story.append(Spacer(1, 12))
    else:
        story.append(Paragraph("Study Compilation", styles['Title']))
        story.append(Spacer(1, 12))
    
    for item in items:
        story.append(Paragraph(item['title'], styles['Heading2']))
        story.append(Paragraph(item['summary'], styles['Normal']))
        bullets = [ListItem(Paragraph(b, styles['Normal']), bulletColor='black') for b in item['bullets']]
        story.append(ListFlowable(bullets, bulletType='bullet'))
        citation = item['citation_apa'] if citation_format == "apa" else item['citation_harvard']
        story.append(Paragraph(citation, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # References
    story.append(Paragraph("References", styles['Heading1']))
    for item in items:
        citation = item['citation_apa'] if citation_format == "apa" else item['citation_harvard']
        story.append(Paragraph(citation, styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_docx(items, atn=None, citation_format="apa"):
    doc = Document()
    
    if atn:
        doc.add_heading(f"Assessment Task: {atn}", 0)
    else:
        doc.add_heading("Study Compilation", 0)
    
    for item in items:
        doc.add_heading(item['title'], 1)
        doc.add_paragraph(item['summary'])
        for bullet in item['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')
        citation = item['citation_apa'] if citation_format == "apa" else item['citation_harvard']
        doc.add_paragraph(citation)
    
    doc.add_heading("References", 1)
    for item in items:
        citation = item['citation_apa'] if citation_format == "apa" else item['citation_harvard']
        doc.add_paragraph(citation)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()