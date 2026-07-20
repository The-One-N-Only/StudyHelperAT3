# -*- coding: utf-8 -*-
import subprocess, os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = r'C:\Users\ramni\OneDrive\Desktop\Software project\StudyHelperAT3'

def set_cell_shading(cell, color='1A1A2E'):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = size
    if alignment is not None: p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2D2D44')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table

def add_image_placeholder(doc, description):
    p = doc.add_paragraph()
    run = p.add_run(f'[IMAGE PLACEHOLDER: {description}]')
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(180, 40, 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def get_git_log():
    try:
        result = subprocess.run(['git', 'log', '--date=short', '--format=%h|%ad|%s', '-100'], capture_output=True, text=True, cwd=PROJECT_DIR)
        lines = result.stdout.strip().split(chr(10))
        return [(p[0], p[1], p[2]) for line in lines if line.count('|') >= 2 for p in [line.split('|', 2)]]
    except Exception:
        return []

print('Helper functions loaded.')
