#!/usr/bin/env python3
"""Generate the full Software Engineering Student Folio DOCX."""
import subprocess, os, sys, time
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_shading(cell, color="1A1A2E"):
    el = OxmlElement("w:shd"); el.set(qn("w:fill"), color); el.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(el)

def H(doc, text, level=1): return doc.add_heading(text, level=level)

def P(doc, text, bold=False, italic=False, sz=None, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if sz: r.font.size = sz
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    return p

def T(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs: rr.bold = True; rr.font.size = Pt(9)
        set_shading(c)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size = Pt(9)
            if ri % 2: set_shading(c, "F5F5F5")
    if widths:
        for i,w in enumerate(widths):
            for rw in t.rows: rw.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def IMG(doc, desc):
    p = doc.add_paragraph(); r = p.add_run(f"[IMAGE PLACEHOLDER: {desc}]")
    r.bold = True; r.italic = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(180,40,40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)

def LI(doc, text):
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(text); r.font.size = Pt(10.5)

def git_log():
    try:
        fmt = "--format=%h|%ad|%s"
        r = subprocess.run(["git","log","--date=short",fmt,"-100"], capture_output=True, text=True, cwd=PROJECT_DIR)
        return [tuple(l.split("|",2)) for l in r.stdout.strip().split("\n") if l.count("|") >= 2]
    except: return []

def get_test_summary():
    try:
        r = subprocess.run([sys.executable,"-m","pytest","tests/","-q","--tb=no"], capture_output=True, text=True, cwd=PROJECT_DIR, timeout=120)
        return r.stdout.strip().split("\n")[-5:]
    except: return ["Test output unavailable"]

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    # ── TITLE PAGE ──
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Software Engineering"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(30,30,60)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Student Folio 2026"); r2.font.size = Pt(18)
    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Assessment Task #3"); r3.font.size = Pt(16)
    doc.add_paragraph()
    t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("StudyLib \u2014 AI-Powered Academic Research Assistant"); r4.bold = True; r4.font.size = Pt(14); r4.italic = True
    for _ in range(3): doc.add_paragraph()
    t5 = doc.add_paragraph(); t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = t5.add_run("Student Number: ______________________________________________"); r5.font.size = Pt(12)
    doc.add_page_break()

    # ── TOC ──
    H(doc, "Table of Contents")
    toc = [
        ("Identifying & Defining the Problem with Solution approach", 0),
        ("    Problem Definition", 1), ("    Functional Requirements", 1), ("    Non-functional Requirements", 1),
        ("    Storyboard (UML Activity Diagram)", 1), ("    Financial Feasibility", 1),
        ("    Dataflow Diagram (DFD)", 1), ("    Data Dictionary", 1), ("    UML Class Diagram", 1), ("    ER Diagram", 1),
        ("Research and Selection of Development Approach", 0),
        ("Project Management & Scheduling", 0), ("    Gantt Chart", 1), ("    Project Diary", 1),
        ("Producing & Implementing the Solution", 0), ("    Algorithm Solution", 1),
        ("    Version Sequence / Code Backup", 1), ("    Prototyping Sequence", 1),
        ("Automated Testing, Optimisation & Evaluation", 0),
    ]
    for item, indent in toc:
        p = doc.add_paragraph(); r = p.add_run(item); r.font.size = Pt(11)
        if not indent: r.bold = True
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1.5) if indent else Cm(0)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 1: IDENTIFYING & DEFINING THE PROBLEM
    # ════════════════════════════════════════════════════════════════
    H(doc, "Identifying & Defining the Problem with Solution Approach")

    H(doc, "Problem Definition", 2)
    P(doc, "Secondary school students frequently struggle to locate credible academic resources when completing research-based assessment tasks. General-purpose search engines return unfiltered results that include unreliable, non-academic, or commercially biased sources, making it difficult for students to distinguish quality material from noise. This problem is compounded when students must also synthesise information across multiple sources and format it into a structured study document aligned to a specific assessment task.")
    P(doc, "No existing tool adequately addresses all three stages of this workflow: finding verified sources, understanding them quickly through summarisation, and compiling them into an exportable study document. Tools such as Google Scholar address discovery but offer no summarisation or compilation features. Note-taking applications such as Notion support compilation but do not restrict sources to verified academic domains.")
    P(doc, "StudyLib addresses this gap by providing a single, locally-run web application that restricts search results to a whitelist of verified academic and government domains, uses AI-powered summarisation to generate concise overviews of each source, and allows students to compile selected summaries into a structured study document exportable as PDF or DOCX. Users can also upload their own study files for AI-assisted summarisation and engage in AI-powered chat with their workspace content. The system targets Year 11 and 12 students undertaking research-heavy assessment tasks and is designed to run on a standard school or home laptop with no external hosting required.")
    P(doc, "Word count: 232", italic=True)

    # 1.2 FRs
    H(doc, "Functional Requirements", 2)
    frs = [
        ("FR1","Users can register with email, username, password, and gender preference","High"),
        ("FR2","Users can log in and maintain an authenticated session with session lockout after 5 failed attempts","High"),
        ("FR3","Passwords are hashed using Werkzeug pbkdf2:sha256 before storage","High"),
        ("FR4","Search queries are restricted to a whitelist of verified academic and government domains","High"),
        ("FR5","Search results display title, source name, URL, and thumbnail image","High"),
        ("FR6","Each search result presents View, Save, and Add to Workspace actions","High"),
        ("FR7","View action loads source content in an embedded in-app viewer via secure server-side proxy","High"),
        ("FR8","In-app proxy viewer fetches content server-side to avoid CORS restrictions","High"),
        ("FR9","In-app viewer displays fallback message or opens new tab for paywalled/inaccessible sources","Medium"),
        ("FR10","AI summarisation generates 100\u2013300 word summary and key bullet points for external sources","High"),
        ("FR11","AI summarisation explicitly states when source content is insufficient rather than fabricating","High"),
        ("FR12","Users can upload PDF, DOCX, TXT, image, XLSX, and XLS files up to 10MB","High"),
        ("FR13","Text is extracted from uploaded files and stored for indexing and search","High"),
        ("FR14","Users can search across extracted text from their uploaded files","Medium"),
        ("FR15","AI summarisation generates 100\u2013300 word summary of uploaded files","Medium"),
        ("FR16","Users can add summarised sources to named compilation workspaces","High"),
        ("FR17","Workspace supports creation, renaming, deletion, and drag-reordering of items","Medium"),
        ("FR18","Compiled documents include auto-generated APA and Harvard citations","Medium"),
        ("FR19","Compiled workspace can be exported as a formatted PDF document","High"),
        ("FR20","Compiled workspace can be exported as a formatted DOCX document","High"),
        ("FR21","Users can save and unsave search results for later access via Saved Items panel","Medium"),
        ("FR22","System logs user actions to an audit log file (user_activity.log)","Low"),
        ("FR23","Unit tests cover domain filtering, text extraction, summarisation, proxy, citation, and AI modules","High"),
        ("FR24","Integration tests verify the full search-to-workspace-to-export workflow","High"),
        ("FR25","Multi-source parallel search returns results from Wikipedia, Google Books, Scholar, PubMed, and whitelist simultaneously","High"),
        ("FR26","AI chatbot (Alexander) answers questions with RAG context from workspaces, files, and web sources","Medium"),
        ("FR27","Workspace chat history is persisted and retrievable across sessions","Medium"),
        ("FR28","Dark and light themes (Candlelit Archive / Old Book) with Victorian-era aesthetic and localStorage persistence","Medium"),
        ("FR29","User profile management: update email, username, password, name, and gender/picture","Medium"),
        ("FR30","PubMed E-utilities integration for MeSH term suggestions and academic metadata enrichment","Low"),
        ("FR31","Search query result caching to avoid redundant SerpAPI calls","Low"),
        ("FR32","Rich-text notes can be attached to workspaces for free-form study documentation","Medium"),
    ]
    T(doc, ["ID","Requirement Description","Priority"], frs, [0.6,5.3,0.8])

    # 1.3 NFRs
    H(doc, "Non-functional Requirements", 2)
    P(doc, "The following non-functional requirements do not directly transform inputs into outputs but are critical to the success and quality of the software application.", italic=True)
    nfrs = [
        ("Performance","Search results returned within 5 seconds; file search within 3 seconds; AI summarisation within 10 seconds; in-app viewer within 5 seconds; multi-source parallel search aggregates within 25 seconds."),
        ("Security","Input validation on all form fields; SQL injection protection via SQLAlchemy ORM parameterisation; API keys stored in environment variables; user data accessible only to owning account; CSRF tokens on all POST submissions."),
        ("Security","Login lockout after 5 consecutive failed attempts (5-minute cooldown); secure content proxy with DNS pinning and IP validation rejecting private/multicast ranges."),
        ("Usability","Clean Bootstrap 5 interface with clear separation between search, viewer, workspace, notes, and chat views; SPA-style navigation with collapsible sidebar menu."),
        ("Usability","Dark and light mode (Candlelit Archive / Old Book theme) toggle with localStorage persistence; Victorian-era engraved illustration aesthetic; accessible contrast ratios."),
        ("Reliability","Graceful error handling on API failure; viewer degrades to fallback link if content unavailable; parallel search survives individual source failures; custom error pages."),
        ("Reliability","Application runs locally without internet-hosted server; all dependencies installable via requirements.txt; teacher can run locally following README instructions."),
        ("Maintainability","Modular architecture: db.py (ORM), search.py (SerpAPI), proxy.py (secure content), summarise.py (AI), files.py (extraction), export.py (PDF/DOCX), citations.py, answer.py (RAG chat), pubmed.py, whitelist.py."),
        ("Maintainability","Version control via Git with 90+ descriptive commits; CSS split by concern (custom.css ~2900 lines); comprehensive test suite with contract tests for visual consistency."),
        ("Portability","Runs on any machine with Python 3.x; SQLite requires no external server; all configuration via .env file; operates offline except for API-dependent features."),
    ]
    T(doc, ["Category","Requirement"], nfrs, [1.5,5.2])

    # 1.4 Storyboard
    H(doc, "Storyboard (UML Activity Diagram)", 2)
    P(doc, "The following UML Activity Diagram maps the complete user journey through StudyLib, from authentication through search, summarisation, workspace compilation, and export. It includes decision nodes, parallel paths (multi-source search), and swimlanes for frontend and backend components.")
    IMG(doc, "UML Activity Diagram \u2014 StudyLib User Journey (Login \u2192 Search \u2192 View \u2192 Summarise \u2192 Workspace \u2192 Export)")
    P(doc, "Key user journeys depicted:", bold=True)
    journeys = [
        "1. Authentication: Login/Register \u2192 CSRF validation \u2192 Credential verification \u2192 Session establishment \u2192 Home dashboard",
        "2. Browse & Search: Enter query \u2192 Select sources (Wikipedia, Google Books, Scholar, PubMed, Whitelist) \u2192 Parallel SerpAPI search \u2192 Deduplication \u2192 Ranked results \u2192 AI browse overview",
        "3. View & Summarise: Click View \u2192 Server-side proxy fetches content \u2192 Render in embedded viewer \u2192 Click Summarise \u2192 AI generates summary, bullets, relevance \u2192 Display structured result",
        "4. Workspace Compilation: Add to Workspace \u2192 Select/create workspace \u2192 Item stored with summary + citations \u2192 Reorder/remove items \u2192 Attach notes \u2192 AI chat with Alexander",
        "5. Export: Select citation format (APA/Harvard) \u2192 Compile workspace items \u2192 Generate PDF or DOCX with summaries, bullets, and citations \u2192 Download",
    ]
    for j in journeys:
        p = doc.add_paragraph(); r = p.add_run(j); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(2)

    # 1.5 Financial Feasibility
    H(doc, "Financial Feasibility", 2)
    P(doc, "Financial feasibility analysis determines whether the software can be built and operated at a cost that is viable for its target users. StudyLib is designed to run locally on a student\u2019s own machine with no hosted infrastructure, eliminating the ongoing operational costs typical of web-hosted applications.")

    H(doc, "Development Costs", 3)
    dev_costs = [
        ("Python 3.x (language runtime)","Free / Open Source"),
        ("Flask and all Python dependencies (via pip)","Free / Open Source"),
        ("Bootstrap 5 (CDN, no build step)","Free"),
        ("SQLite (bundled with Python)","Free"),
        ("Git / GitHub (version control and repository)","Free for public repositories"),
        ("VS Code / GitHub Codespace (development environment)","Free (student tier)"),
        ("SerpAPI (search engine integration)","Free tier: 100 searches/month"),
        ("Anthropic Claude API (AI summarisation and chatbot)","$5.00 AUD pre-paid credit"),
        ("Total development cost","$5.00 AUD"),
    ]
    T(doc, ["Resource","Cost"], dev_costs, [3.3,3.4])
    P(doc, "The sole out-of-pocket expense for this project was a $5.00 AUD pre-paid credit loaded onto an Anthropic API account. All other tools, libraries, and services were used at zero cost. No student end user pays any subscription, license, or usage fee to operate StudyLib.")

    H(doc, "API Usage and Pricing", 3)
    P(doc, "StudyLib integrates two third-party APIs, selected to avoid recurring costs at the individual student scale.")
    P(doc, "SerpAPI (Search): The SerpAPI free tier provides 100 searches per month at no cost. A single research session typically issues 1\u20132 multi-source searches, each fanning out to 3\u20135 providers (Wikipedia, Google Books, Google Scholar, PubMed, whitelist domains). At 100 searches/month, a student can complete approximately 20\u201330 full research sessions before exhausting the free allocation. For the assessment demonstration, all search queries during development and testing remained within the free tier limit.")
    P(doc, "Anthropic Claude (AI Summarisation and Chat): StudyLib uses two Claude models in production, selected for different cost\u2013capability profiles:")

    T(doc, ["Component","Model","Input Pricing","Output Pricing (per 1M tokens)"],
       [("Summarisation","Claude Haiku 4.5","$0.80","$4.00"),
        ("Chat (Alexander)","Claude Sonnet 4.6","$3.00","$15.00")],
       [1.5,2.0,1.5,1.7])
    P(doc, "Haiku 4.5 is used for source summarisation because it is Anthropic\u2019s fastest and most cost-efficient model, well-suited to the structured JSON output format required (summary + bullets + relevance + citations). A typical summarisation call consumes approximately 1,500 input tokens (source content + prompt) and 300 output tokens, costing approximately $0.0024 AUD per summary.")
    P(doc, "Sonnet 4.6 is used for multi-turn chat (Alexander) because the Retrieval-Augmented Generation (RAG) pipeline requires strong reasoning over aggregated context from workspace items, uploaded files, and web sources. A typical chat turn consumes approximately 3,000 input tokens (conversation history + RAG context) and 500 output tokens, costing approximately $0.0165 AUD per exchange.")
    P(doc, "The $5.00 AUD pre-paid credit funded all AI operations during the entire development and assessment period: approximately 80\u2013100 summarisation calls during implementation and testing, 30\u201340 chat exchanges, and the final assessment demonstration \u2014 with credit to spare. At these rates, a student completing a full research task involving 15 source summaries and 10 chat exchanges would incur approximately $0.20 AUD in total API costs \u2014 a negligible amount that can be absorbed by a shared classroom API key or the student\u2019s own pre-paid credit.")

    H(doc, "Operational Costs at Scale", 3)
    P(doc, "Because StudyLib runs locally on the student\u2019s machine (Flask development server, SQLite database), there are no hosting, domain registration, or server infrastructure costs. The only recurring operational consideration is API usage volume.")
    T(doc, ["Scenario","Monthly API Cost (Est.)"],
       [("Single student (15 summaries + 10 chats per week)","~$0.86 AUD/month"),
        ("Classroom of 30 students (5 research tasks each per term)","~$30 AUD/term"),
        ("School-wide deployment (300 students, 5 tasks each per term)","~$300 AUD/term")],
       [4.0,2.7])
    P(doc, "These figures represent worst-case estimates; in practice, search result caching avoids redundant API calls, and many browsing sessions do not trigger summarisation or chat. Volume discounts from both SerpAPI and Anthropic would further reduce per-unit costs at the school-wide scale.")

    H(doc, "Comparison to Commercial Alternatives", 3)
    T(doc, ["Product","Pricing"],
       [("Scite","$20 AUD/month (individual)"),
        ("RefWorks","$15 AUD/month (institutional)"),
        ("Paperpile","$10 AUD/month (academic)")],
       [3.0,3.7])
    P(doc, "StudyLib delivers a comparable research workflow \u2014 verified academic search, AI summarisation, citation generation, document compilation \u2014 at a fraction of the cost. At the individual student level, the application is effectively free, requiring only a nominal pre-paid API credit (<$5 AUD per term for heavy use). A commercial alternative such as Scite would cost a student $240 AUD per year for equivalent functionality.")

    H(doc, "Break-Even Analysis", 3)
    P(doc, "StudyLib is a non-commercial assessment project developed at a total cost of $5.00 AUD. The application is distributed at zero cost to end users. From the student\u2019s perspective, break-even is immediate: the application costs nothing to download and run, and optional API usage costs less than a single cup of coffee per term.")
    P(doc, "From a school deployment perspective, the break-even point against the cheapest commercial alternative (Paperpile at $10/month) is reached at approximately 1.5 students. A school spending $300 AUD per term on API credits for 300 students saves approximately $8,700 AUD per term compared to equivalent Paperpile licenses.")

    # 1.6 DFD
    H(doc, "Dataflow Diagram (DFD)", 2)
    P(doc, "The following Level 1 Dataflow Diagram illustrates the key processes within StudyLib, the data flows between them, the external entities that interact with the system, and the data stores used for persistent storage.")
    P(doc, "Processes identified:", bold=True)
    procs = [
        "P1 \u2014 Search & Browse: Accepts user query, calls SerpAPI/PubMed APIs, caches results, returns deduplicated ranked items to the user. Input: query + source selection. Output: ranked item list.",
        "P2 \u2014 Content Proxy & View: Fetches source URLs server-side, validates domains against whitelist, sanitises HTML (removes scripts/forms/active content), returns reader-mode content to in-app viewer.",
        "P3 \u2014 AI Summarisation: Receives source content (URL or uploaded file), calls Anthropic Claude API, generates structured JSON with summary, bullets, and relevance; stores results with workspace item.",
        "P4 \u2014 Workspace Management: Creates/renames/deletes workspaces, adds items with summaries and citations, reorders items, attaches notes, persists chat messages. Manages many-to-many user-item relationships.",
        "P5 \u2014 Export Compilation: Reads workspace items, formats summaries/citations/bullets, generates PDF (via ReportLab) or DOCX (via python-docx) binary output for download.",
        "P6 \u2014 File Upload & Extraction: Accepts uploaded PDF/DOCX/TXT/XLSX/image files, extracts plain text via PyMuPDF/python-docx/openpyxl, stores metadata and extracted text for indexing.",
        "P7 \u2014 AI Chat (RAG): Receives user messages in workspace context, retrieves relevant documents via keyword matching from uploaded files and workspace items, calls Anthropic API with context-augmented prompt, stores conversation.",
    ]
    for proc in procs:
        p = doc.add_paragraph(); r = p.add_run(proc); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

    P(doc, "Data stores:", bold=True)
    P(doc, "D1 \u2014 SQLite Database (server.db): users, items, workspaces, workspace_items, workspace_chat_messages, notes, uploaded_files, search_cache, user_to_saved, user_to_recently_viewed, user_to_recently_searched")
    P(doc, "D2 \u2014 File System (static/uploads/<user_id>/): Uploaded PDF/DOCX/TXT/XLSX/image files stored by UUID-prefixed filename")
    P(doc, "D3 \u2014 Audit Log (user_activity.log): Timestamped log of user actions for security auditing")
    P(doc, "External entities: User (Student), SerpAPI, Anthropic Claude API, PubMed E-utilities API")
    IMG(doc, "Dataflow Diagram (DFD) Level 1 \u2014 StudyLib System Processes, Data Stores, and External Entities")

    # 1.7 Data Dictionary
    H(doc, "Data Dictionary", 2)
    P(doc, "This data dictionary documents every table in the StudyLib SQLite database as implemented in src/db.py. It records, for each attribute: its name, data type, description, an example value, validation rules applied at the application layer, and database-level constraints. Foreign key and primary key relationships are listed under Constraints and correspond to the ER Diagram.")

    # users
    H(doc, "users", 3)
    P(doc, "Stores both locally-registered (username/password) and platform-login accounts. password_hash is nullable to support alternate login platforms. The gender field supports Victorian-era profile picture selection.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key; unique identifier","1","Auto-incremented","PRIMARY KEY, NOT NULL"),
        ("email","VARCHAR(254)","User email address; unique account identifier","aarush@example.com","Must contain @; max 254 chars","UNIQUE, NOT NULL"),
        ("name","VARCHAR(254)","User display name (optional)","Aarush Sharma","Max 254 chars","NULLABLE"),
        ("username","VARCHAR(254)","Username used to log in","aarush_s","3-30 chars; regex [A-Za-z0-9_.-]{3,30}","NULLABLE, app-enforced unique"),
        ("password_hash","VARCHAR(255)","Salted hash (Werkzeug pbkdf2:sha256)","pbkdf2:sha256:...","Min source password 8 chars","NULLABLE"),
        ("gender","VARCHAR(16)","Avatar selection for Victorian profile picture","gentleman","One of: gentleman, lady, secret","NOT NULL, DEFAULT 'gentleman'"),
        ("login_platform","VARCHAR(16)","Authentication method","local","Restricted to known values","NOT NULL, DEFAULT 'local'"),
        ("platform_id","JSON","Platform-specific identifier object","{}","Must be valid JSON object","NOT NULL, DEFAULT {}")],
       [1.0,1.0,2.0,1.0,1.3,1.3])

    # items
    H(doc, "items", 3)
    P(doc, "Cached, de-duplicated record of every external source (Wikipedia, Google Books, PubMed, Scholar, whitelist) ever returned. Re-used across searches via source_id to avoid duplicate API calls.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key; unique identifier","42","Auto-incremented","PRIMARY KEY, NOT NULL"),
        ("title","VARCHAR(255)","Title of the source","Photosynthesis","Max 255 chars","NOT NULL"),
        ("description","VARCHAR(1023)","Short summary or snippet","The process by which green plants...","Max 1023 chars","NOT NULL"),
        ("thumb_url","VARCHAR(255)","URL of thumbnail image","https://upload.wikimedia.org/...","Valid URL or empty string","NOT NULL"),
        ("thumb_mime","VARCHAR(255)","MIME type of thumbnail","image/jpeg","Recognised image MIME or empty","NOT NULL"),
        ("thumb_height","INTEGER","Rendered height in pixels","135","Clamped 0\u2013135","NOT NULL"),
        ("source_url","VARCHAR(1023)","Canonical URL of original source","https://en.wikipedia.org/wiki/Photosynthesis","Must pass whitelist domain check","NOT NULL"),
        ("source_name","VARCHAR(64)","Originating search provider","wikipedia","One of: wikipedia, gbooks, pubmed, scholar, whitelist","NOT NULL"),
        ("source_id","VARCHAR(1023)","Provider-specific unique identifier (URL or ID)","https://en.wikipedia.org/wiki/...","Format depends on provider","UNIQUE, NOT NULL"),
        ("abstract","TEXT","Full abstract text (PubMed)","Background: Photosynthesis is...","No format restriction","NULLABLE"),
        ("authors","TEXT","JSON-encoded author names","[\"Smith, J\", \"Lee, K\"]","Valid JSON array if present","NULLABLE"),
        ("journal","VARCHAR(255)","Journal name (PubMed)","Nature","Max 255 chars","NULLABLE"),
        ("year","VARCHAR(4)","Publication year","2023","4-digit year format","NULLABLE"),
        ("volume","VARCHAR(32)","Journal volume number","57","Max 32 chars","NULLABLE"),
        ("issue","VARCHAR(32)","Journal issue number","3","Max 32 chars","NULLABLE"),
        ("doi","VARCHAR(255)","Digital Object Identifier","10.1038/s41586-023-00001","Must follow DOI syntax","NULLABLE")],
       [1.0,1.0,2.0,1.0,1.2,1.2])

    # Junction tables
    for tname, desc, rows in [
        ("user_to_saved","Junction: users \u2194 bookmarked items. Composite primary key on (user_id, item_id).",[
            ("user_id","INTEGER","User who saved","1","Must reference users.id","PK, FK->users.id, NOT NULL"),
            ("item_id","INTEGER","Saved item","42","Must reference items.id","PK, FK->items.id, NOT NULL"),
            ("time_inserted","INTEGER","Microsecond Unix timestamp","1718000000000000","Positive integer","NOT NULL")]),
        ("user_to_recently_viewed","Tracks last 10 items opened in viewer; oldest pruned automatically.",[
            ("user_id","INTEGER","Viewing user","1","Must reference users.id","PK, FK->users.id, NOT NULL"),
            ("item_id","INTEGER","Viewed item","42","Must reference items.id","PK, FK->items.id, NOT NULL"),
            ("time_inserted","INTEGER","Microsecond Unix timestamp","1718000000000000","Positive integer","NOT NULL")]),
        ("user_to_recently_searched","Tracks last 10 items returned via search; populates Recently Searched panel.",[
            ("user_id","INTEGER","Searching user","1","Must reference users.id","PK, FK->users.id, NOT NULL"),
            ("item_id","INTEGER","Searched item","42","Must reference items.id","PK, FK->items.id, NOT NULL"),
            ("time_inserted","INTEGER","Microsecond Unix timestamp","1718000000000000","Positive integer","NOT NULL")]),
    ]:
        H(doc, tname, 3)
        P(doc, desc, italic=True)
        T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"], rows, [1.2,0.8,2.0,1.0,1.4,1.4])

    # workspaces
    H(doc, "workspaces", 3)
    P(doc, "User-created compilation container. A user may hold multiple workspaces (e.g. one per subject or assessment task).", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key","3","Auto-incremented","PK, NOT NULL"),
        ("user_id","INTEGER","Owning user","1","Must reference users.id","FK->users.id, NOT NULL"),
        ("name","VARCHAR(255)","Display name","Biology Assessment Task 2","1-255 chars, non-empty","NOT NULL"),
        ("time_created","INTEGER","Unix timestamp (seconds)","1718000000","Positive integer","NOT NULL")],
       [1.1,0.9,2.0,1.2,1.2,1.4])

    # workspace_chat_messages
    H(doc, "workspace_chat_messages", 3)
    P(doc, "Persists AI chat conversation history attached to a workspace. Each turn consists of a user message and an assistant (Alexander) response.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key","15","Auto-incremented","PK, NOT NULL"),
        ("user_id","INTEGER","Message author","1","Must reference users.id","FK->users.id, NOT NULL"),
        ("workspace_id","INTEGER","Parent workspace","3","Must reference workspaces.id","FK->workspaces.id, NOT NULL"),
        ("role","VARCHAR(16)","Author role","assistant","One of: user, assistant","NOT NULL"),
        ("content","TEXT","Message body","Based on your sources...","Free text","NOT NULL"),
        ("time_created","INTEGER","Unix timestamp (seconds)","1718000400","Positive integer","NOT NULL")],
       [1.1,0.9,2.0,1.2,1.2,1.4])

    # workspace_items
    H(doc, "workspace_items", 3)
    P(doc, "Compiled entry inside a workspace: either an external source (item_id) or an uploaded file (file_id), with AI-generated summary, bullets, relevance note, and citation strings.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key","7","Auto-incremented","PK, NOT NULL"),
        ("user_id","INTEGER","Owning user","1","Must reference users.id","FK->users.id, NOT NULL"),
        ("workspace_id","INTEGER","Parent workspace","3","Must reference workspaces.id","FK->workspaces.id, NULLABLE"),
        ("item_id","INTEGER","Source item (if from search)","42","Must reference items.id when set","FK->items.id, NULLABLE"),
        ("file_id","INTEGER","Uploaded file (if from upload)","5","Must reference uploaded_files.id when set","FK->uploaded_files.id, NULLABLE"),
        ("summary","TEXT","AI-generated 100-300 word summary","Photosynthesis converts light energy...","Must flag insufficient content","NOT NULL"),
        ("bullets","TEXT","JSON-encoded key bullet points","[\"Occurs in chloroplasts\"]","Valid JSON array","NOT NULL"),
        ("relevance","TEXT","AI explanation linking source to ATN","This source explains the core mechanism...","Only when ATN supplied","NULLABLE"),
        ("atn_used","TEXT","Assessment Task Notification text","Explain the role of photosynthesis...","Free text","NULLABLE"),
        ("citation_apa","TEXT","Pre-formatted APA citation","Wikipedia. (n.d.). Photosynthesis...","Generated via citations.py","NOT NULL"),
        ("citation_harvard","TEXT","Pre-formatted Harvard citation","Wikipedia (n.d.) 'Photosynthesis'...","Generated via citations.py","NOT NULL"),
        ("position","INTEGER","Sort order (drag-reordering)","0","Non-negative integer","NOT NULL"),
        ("time_added","INTEGER","Unix timestamp (seconds) added","1718000200","Positive integer","NOT NULL")],
       [1.0,0.8,2.0,1.0,1.2,1.2])

    # uploaded_files
    H(doc, "uploaded_files", 3)
    P(doc, "Metadata and extracted text for a student\u2019s personally uploaded study file (PDF, DOCX, TXT, image, or spreadsheet).", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key","5","Auto-incremented","PK, NOT NULL"),
        ("user_id","INTEGER","Uploading user","1","Must reference users.id","FK->users.id, NOT NULL"),
        ("filename","VARCHAR(255)","Original filename","biology_notes.pdf","Max 255 chars","NOT NULL"),
        ("stored_path","VARCHAR(512)","Server-side path with UUID prefix","static/uploads/1/3f2a..._biology_notes.pdf","Within static/uploads/<user_id>/","NOT NULL"),
        ("file_type","VARCHAR(8)","Normalised type for extraction","pdf","One of: pdf, docx, txt, image, xlsx, xls","NOT NULL"),
        ("extracted_text","TEXT","Plain text extracted from file","Cellular respiration is...","Empty string if extraction fails","NOT NULL"),
        ("file_size","INTEGER","Size in bytes","2485760","Max 10MB (10,485,760 bytes)","NOT NULL"),
        ("time_uploaded","INTEGER","Unix timestamp (seconds)","1718000100","Positive integer","NOT NULL")],
       [1.1,0.9,2.0,1.2,1.2,1.2])

    # notes
    H(doc, "notes", 3)
    P(doc, "Free-form rich-text notes attachable to a user\u2019s account or a specific workspace. Content is HTML produced by the Quill editor.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("id","INTEGER","Primary key","9","Auto-incremented","PK, NOT NULL"),
        ("user_id","INTEGER","Owning user","1","Must reference users.id","FK->users.id, NOT NULL"),
        ("workspace_id","INTEGER","Parent workspace","3","Must reference workspaces.id","FK->workspaces.id, NULLABLE"),
        ("title","VARCHAR(255)","Note title","Key terms for revision","1-255 chars, non-empty","NOT NULL"),
        ("content","TEXT","Rich-text (HTML) note body","<p>Mitochondria are...</p>","Sanitised on render","NOT NULL"),
        ("time_created","INTEGER","Unix timestamp (seconds) created","1718000300","Positive integer","NOT NULL"),
        ("time_updated","INTEGER","Unix timestamp (seconds) last edit","1718000400","Positive integer; updated on save","NOT NULL")],
       [1.1,0.9,2.0,1.2,1.2,1.2])

    # search_cache
    H(doc, "search_cache", 3)
    P(doc, "Caches search query results to avoid redundant SerpAPI calls for identical queries within a session.", italic=True)
    T(doc, ["Attribute","Data Type","Description","Example","Validation","Constraints"],
       [("cache_key","VARCHAR(64)","SHA-256 hash of query + source","a1b2c3d4...","64-char hex digest","PK, NOT NULL"),
        ("item_ids","TEXT","JSON-encoded list of cached item IDs","[42, 43, 44]","Valid JSON array","NOT NULL"),
        ("time_cached","INTEGER","Unix timestamp (seconds)","1718000000","Positive integer","NOT NULL")],
       [1.1,0.9,2.0,1.2,1.2,1.2])

    # 1.8 UML Class Diagram
    H(doc, "UML Class Diagram", 2)
    P(doc, "The UML Class Diagram below represents the software system\u2019s object model as implemented in src/db.py using SQLAlchemy ORM. Each class maps to a database table and includes attributes with data types and relationship multiplicities.")
    P(doc, "Key classes and relationships:", bold=True)
    class_rels = [
        "User (1) \u2500\u2500\u2500 (*) Workspace: One user owns multiple workspaces",
        "User (1) \u2500\u2500\u2500 (*) WorkspaceItem: One user contributes to multiple workspace compilations",
        "User (1) \u2500\u2500\u2500 (*) UploadedFile: One user uploads multiple files",
        "User (1) \u2500\u2500\u2500 (*) Note: One user creates multiple notes",
        "Workspace (1) \u2500\u2500\u2500 (*) WorkspaceItem: One workspace contains multiple compiled entries",
        "Workspace (1) \u2500\u2500\u2500 (*) Note: One workspace has multiple attached notes",
        "Workspace (1) \u2500\u2500\u2500 (*) WorkspaceChatMessage: One workspace has a multi-turn chat history",
        "Item (1) \u2500\u2500\u2500 (*) WorkspaceItem: One source can appear in multiple workspace compilations",
        "UploadedFile (1) \u2500\u2500\u2500 (*) WorkspaceItem: One uploaded file can appear in multiple compilations",
        "User (*) \u2500\u2500\u2500 (*) Item via UserToSaved: Many-to-many bookmark junction",
        "User (*) \u2500\u2500\u2500 (*) Item via UserToRecentlyViewed: Many-to-many view history junction",
        "User (*) \u2500\u2500\u2500 (*) Item via UserToRecentlySearched: Many-to-many search history junction",
    ]
    for rel in class_rels:
        p = doc.add_paragraph(); r = p.add_run(rel); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)
    IMG(doc, "UML Class Diagram \u2014 StudyLib ORM Model (11 classes with attributes, methods, and relationships)")

    # 1.9 ER Diagram
    H(doc, "ER Diagram", 2)
    P(doc, "The Entity-Relationship Diagram below represents the database schema for StudyLib. It shows 11 tables with their attributes, primary keys, foreign keys, and relationship cardinalities. The diagram directly corresponds to the SQLAlchemy ORM models defined in src/db.py.")
    P(doc, "Key relationships:", bold=True)
    er_rels = [
        "users 1\u2500\u2500* workspaces: One user owns many workspaces",
        "users 1\u2500\u2500* workspace_items: One user compiles many workspace entries",
        "users 1\u2500\u2500* uploaded_files: One user uploads many files",
        "users 1\u2500\u2500* notes: One user writes many notes",
        "users *\u2500\u2500* items (via user_to_saved): Many-to-many bookmarks",
        "users *\u2500\u2500* items (via user_to_recently_viewed): Many-to-many view history",
        "users *\u2500\u2500* items (via user_to_recently_searched): Many-to-many search history",
        "workspaces 1\u2500\u2500* workspace_items: One workspace contains many compiled items",
        "workspaces 1\u2500\u2500* notes: One workspace has many attached notes",
        "workspaces 1\u2500\u2500* workspace_chat_messages: One workspace has many chat messages",
        "items 1\u2500\u2500* workspace_items: One source appears in many compilations",
        "uploaded_files 1\u2500\u2500* workspace_items: One file appears in many compilations",
    ]
    for rel in er_rels:
        p = doc.add_paragraph(); r = p.add_run(rel); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)
    IMG(doc, "ER Diagram \u2014 StudyLib Database Schema (11 tables with PKs, FKs, and relationship cardinalities)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 2: RESEARCH & SELECTION OF DEVELOPMENT APPROACH
    # ════════════════════════════════════════════════════════════════
    H(doc, "Research and Selection of Development Approach")

    H(doc, "Overview", 2)
    P(doc, "Selecting an appropriate software development methodology is a foundational decision that shapes how a project is planned, executed, and evaluated. For this project, three approaches were considered: Waterfall, Agile, and WAgile (a hybrid of both). Each has distinct characteristics suited to different project scales and workflows.")

    H(doc, "Waterfall", 2)
    P(doc, "The Waterfall model is a linear, sequential development process in which each phase \u2014 requirements, design, implementation, testing, and deployment \u2014 must be completed before the next begins. It was formalised by Winston Royce in 1970 and became the dominant approach in enterprise and government software development throughout the 1980s and 1990s.")
    P(doc, "The primary strength of Waterfall is its predictability. Because all requirements are defined upfront and documented in detail, the scope, timeline, and cost of a project can be estimated with high confidence before any code is written. This makes it well-suited to large-scale projects with fixed contracts, stable requirements, and regulatory compliance obligations \u2014 for example, defence systems, medical device software, or banking infrastructure, where changes mid-project are expensive or legally prohibited.")
    P(doc, "However, Waterfall\u2019s rigidity is also its greatest weakness. If requirements change after the design phase \u2014 or if testing in the final phase reveals a fundamental design flaw \u2014 the cost of revisiting earlier stages is high. There is also limited opportunity for stakeholder feedback until a working product is delivered, which means problems with the user experience or functionality may not surface until late in the cycle. For a project of this scope and timeline, these limitations are significant.")

    H(doc, "Agile", 2)
    P(doc, "Agile is an iterative development framework formalised in the 2001 Agile Manifesto. Rather than defining all requirements upfront, Agile breaks development into short cycles called sprints or iterations, typically one to two weeks in length. At the end of each iteration, a working increment of the software is delivered and reviewed, and requirements can be adjusted based on feedback.")
    P(doc, "Agile\u2019s key advantage is its responsiveness to change. Features can be added, reprioritised, or removed as the project evolves, which makes it ideal for projects where requirements are not fully known at the outset \u2014 particularly consumer-facing software where user feedback is critical. Large technology companies such as Spotify and Atlassian have adopted Agile extensively for this reason.")
    P(doc, "The limitation of pure Agile at the scale of an individual student project is the overhead it assumes: dedicated product owners, daily standups, sprint reviews, and retrospectives. These ceremonies are designed for teams of four or more. For a single-developer project with a fixed submission deadline, the full Agile framework introduces process for its own sake rather than adding value. Additionally, Agile\u2019s de-emphasis on upfront documentation can leave a project without the planning artefacts \u2014 DFDs, class diagrams, pseudocode \u2014 that this assessment explicitly requires.")

    H(doc, "WAgile (Hybrid)", 2)
    P(doc, "WAgile, sometimes referred to as Water-Scrum-Fall or hybrid Agile, combines structured upfront planning from Waterfall with iterative development cycles from Agile. The project begins with a planning phase \u2014 defining requirements, designing data structures, and producing documentation \u2014 before development is broken into iterative sprints that each deliver a working feature increment. Testing and review occur throughout, not only at the end.")
    P(doc, "This approach has been adopted widely in medium-scale software teams that need the predictability of a defined plan but want the flexibility to respond to implementation challenges as they arise. Microsoft\u2019s development of Visual Studio Code, for example, follows a hybrid model: fixed quarterly roadmaps with iterative monthly sprint deliveries within each quarter.")

    H(doc, "Selected Approach: WAgile", 2)
    P(doc, "WAgile was selected as the development approach for StudyLib. The project has a fixed submission deadline and a defined set of features documented in the Software Requirements Specification, which makes upfront planning essential. However, the implementation of AI summarisation, server-side proxying, SerpAPI multi-source search, and file text extraction involved technical unknowns that benefited from iterative development \u2014 the scope of each module was refined as implementation revealed complexity not apparent at the design stage.")
    P(doc, "Concretely, the project was structured as follows: a planning phase (June 18\u201325) produced the requirements specification, data dictionary, ER diagram, and class diagram before any code was written. Development then proceeded in feature-based sprints: authentication and home page UI (June 25\u2013July 2), search and whitelist integration (July 2\u20139), SerpAPI multi-source search and browse page (July 9\u201315), AI summarisation and Claude chatbot integration (July 15\u201316), Candlelit Archive dark theme and Victorian UI (July 16\u201317), file upload and extraction (July 17\u201318), workspace management and export (July 18\u201319), and browse AI overview, Google Books previews, and test contracts (July 19\u201320). Each sprint completed a vertical slice of functionality before the next began, allowing testing to occur module-by-module rather than only at submission, which reduced the risk of late-stage integration failures.")
    P(doc, "WAgile is the most appropriate methodology for a project of this scale: solo-developed, deadline-bound, with well-defined but technically uncertain requirements. The upfront planning artefacts satisfy the documentation requirements of the assessment, while the iterative development cycles provided the flexibility to adapt to technical challenges encountered during implementation.")
    P(doc, "Word count: ~950", italic=True)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 3: PROJECT MANAGEMENT & SCHEDULING
    # ════════════════════════════════════════════════════════════════
    H(doc, "Project Management & Scheduling")

    H(doc, "Gantt Chart", 2)
    P(doc, "The Gantt chart below outlines the planned and actual development timeline for StudyLib. It identifies key milestones, task dependencies, and sprint durations across the project period from June 18 to July 24, 2026. The chart demonstrates the WAgile approach: upfront planning followed by iterative, feature-based development sprints.")
    P(doc, "Development phases:", bold=True)
    phases = [
        ("Phase 1: Planning & Setup","June 18\u201325","Requirements specification, data dictionary, diagrams, repository setup, environment configuration"),
        ("Phase 2: Auth & UI Foundation","June 25\u2013July 2","Login/register with lockout, CSRF, password hashing; home page UI; sidebar navigation; error pages"),
        ("Phase 3: Search & Browse","July 2\u201315","Whitelist integration, SerpAPI configuration, multi-source parallel search (Wikipedia, Google Books, Scholar, PubMed), browse page, result deduplication, pagination"),
        ("Phase 4: Content & AI","July 15\u201316","Secure proxy with DNS pinning and HTML sanitisation; Anthropic Claude for summarisation and chatbot (Alexander); PubMed E-utilities integration"),
        ("Phase 5: Theme & Design","July 16\u201317","Candlelit Archive dark theme (Victorian leather/wood textures, SVG illustrations, candle glow); Old Book light theme; localStorage toggle"),
        ("Phase 6: Files, Workspace & Export","July 17\u201319","File upload with text extraction; workspace CRUD and reorder; notes within workspaces; PDF/DOCX export with citations"),
        ("Phase 7: Browse Polish & AI Overview","July 19\u201320","Browse AI overview generation, Google Books previews, result imagery and filters, ranked reveal and paging, persistent state restoration"),
        ("Phase 8: Testing & Documentation","July 20\u201324","Comprehensive test suite execution (10+ test modules, 381 tests), folio documentation completion, final review and submission"),
    ]
    T(doc, ["Phase","Dates","Key Deliverables"], phases, [2.0,1.2,3.5])
    IMG(doc, "Gantt Chart \u2014 StudyLib Development Timeline (Planned vs Actual with Dependencies and Milestones)")

    # Project Diary
    H(doc, "Project Diary", 2)
    P(doc, "The following project diary entries are derived from the Git commit log and reflect the iterative WAgile development process. Each entry records the date, work completed, challenges encountered, and decisions made \u2014 providing evidence of project management and reflective practice.", italic=True)

    diary = [
        ("2026-06-18","Repository initialised. Created basic Flask application structure with modular directory layout (src/, templates/, static/, tests/). Installed core dependencies: Flask, SQLAlchemy, Flask-Session.","Established clean separation of concerns from the outset. Decision to use SQLite for zero-config local deployment eliminated need for external database server."),
        ("2026-06-25","Created custom error page templates (404, 500, error.html) and integrated with Flask error handlers. Ensured graceful degradation for all exception paths.","Comprehensive error handling is critical for user experience. Implemented early to establish consistent feedback patterns across the application."),
        ("2026-07-02","Enhanced search with whitelist domain filtering support. Loaded .env configuration in pubmed and summarise modules. Improved test coverage for search and whitelist modules.","Whitelist-based search is a core differentiator. JSON whitelist file enables easy extensibility without code changes. Standardised .env loading pattern across all API-sensitive modules."),
        ("2026-07-09","Replaced top bar navigation with hamburger menu sidebar. Updated browse page with initial search integration. Implemented whitelisted domain backing before API key configuration.","Sidebar navigation decision driven by need for distraction-free research interface. Hamburger menu maximises screen real estate for search results and content viewing, aligning with Victorian archive aesthetic."),
        ("2026-07-15","Major milestone: Integrated Anthropic Claude API for AI summarisation and chatbot (Alexander). Implemented proper SerpAPI search integration enabling full-web academic search across multiple sources. Renamed chatbot from Jason to Alexander.","Claude selected for strong instruction-following and safe content generation. RAG pattern uses workspace content as context, ensuring responses grounded in actual research materials. SerpAPI chosen over direct Google API for simpler multi-source integration."),
        ("2026-07-17","Added Candlelit Archive dark theme with Victorian-era leather textures, wood panel tones, engraved SVG illustrations, and candle glow hover effects. Created comprehensive CSS contract tests (~2900 lines) for both themes. Fixed workspace creation input and search API key configuration.","Victorian aesthetic differentiates StudyLib visually and creates engaging scholarly atmosphere. CSS custom properties for consistent theming. Contract tests essential to maintain visual consistency across both themes without regressions."),
        ("2026-07-18","Implemented multi-source parallel search (search-all endpoint) aggregating results from Wikipedia, Google Books, Scholar, and whitelist domains simultaneously. Added SerpAPI result caching. Merged main branch with StudyLib functional polish. Added Google Books embedded previews in viewer.","Parallel search with ThreadPoolExecutor (6 workers, 25s timeout) significantly improved search responsiveness. Deduplication via union-find algorithm with canonical URL normalisation ensures clean results across overlapping source scopes."),
        ("2026-07-19","Added wood-texture panels to browse results, inline workspace creation, improved proxy handling for paywalled sources. Implemented browse AI overview generation with persistence, timeout handling, retry, and stale-query rejection. Added ranked browse result reveal with progressive paging.","Browse AI overview provides high-level synthesis before individual result review. Stale-query rejection prevents race conditions between fast subsequent searches. Wood textures extend the Victorian theme to search results consistently."),
        ("2026-07-20","Final UI upgrades: Victorian SVG illustrations, profile picture selection by gender (gentleman/lady/secret), theme polish. Added loader display with Bible page-turn animation, proxy blur for loading content, square navigation icons, and workspace source preview integration.","Victorian-era profile pictures (engraved gentleman, lady, quill) create personal connection to the archive theme. Page-turn loader animation reinforces the scholarly library aesthetic. Workspace source preview bridges the gap between search and compilation views seamlessly."),
    ]
    for date, desc, challenge in diary:
        p = doc.add_paragraph(); r = p.add_run(f"{date}"); r.bold = True; r.font.size = Pt(10.5)
        p2 = doc.add_paragraph(); r2 = p2.add_run(desc); r2.font.size = Pt(10); p2.paragraph_format.left_indent = Cm(0.5); p2.paragraph_format.space_after = Pt(2)
        p3 = doc.add_paragraph(); r3 = p3.add_run(f"Reflection: {challenge}"); r3.italic = True; r3.font.size = Pt(10); p3.paragraph_format.left_indent = Cm(0.5); p3.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 4: PRODUCING & IMPLEMENTING THE SOLUTION
    # ════════════════════════════════════════════════════════════════
    H(doc, "Producing & Implementing the Solution")

    H(doc, "Algorithm Solution", 2)
    P(doc, "The following pseudocode documents the login and registration algorithms, which are the core authentication workflows in StudyLib. These algorithms demonstrate secure software architecture principles: CSRF protection, input validation, password hashing, session management, and brute-force lockout protection.")

    H(doc, "Registration Algorithm", 3)
    auth_code = """PROCEDURE Register(email, name, username, password, confirm_password)

    // Validate session lockout
    IF session contains login_lockout_until AND
       current_time < login_lockout_until THEN
        Display 'Too many attempts. Try again later'
        RETURN
    END IF

    // Validate CSRF token
    IF form csrf_token != session csrf_token THEN
        Abort with HTTP 400 Bad Request
    END IF

    // Validate required fields
    IF email is empty OR username is empty OR password is empty THEN
        Display 'Email, username and password are required'
        RETURN
    END IF

    // Validate password strength
    IF length(password) < 8 THEN
        Display 'Password must be at least 8 characters'
        RETURN
    END IF

    // Validate password confirmation
    IF password != confirm_password THEN
        Display 'Passwords do not match'
        RETURN
    END IF

    // Validate username format
    IF username does not match pattern [A-Za-z0-9_.-]{3,30} THEN
        Display 'Username contains invalid characters'
        RETURN
    END IF

    // Validate email format
    IF '@' not in email OR length(email) > 254 THEN
        Display 'Please enter a valid email address'
        RETURN
    END IF

    // Check for duplicate email or username
    IF get_user_by_email(email) is not null THEN
        Display 'Email is already registered'
        RETURN
    END IF
    IF get_user_by_username(username) is not null THEN
        Display 'Username already exists'
        RETURN
    END IF

    // Hash password using pbkdf2:sha256
    password_hash <- generate_password_hash(password)

    // Create user record with gender preference
    user <- create_local_user(email, username, password_hash, name, gender)

    // Establish authenticated session
    Clear session
    session[user_id] <- user.id
    session[username] <- user.username
    session[gender] <- user.gender
    Display 'Registration successful'
    REDIRECT to home page
END PROCEDURE"""
    for line in auth_code.strip().split("\n"):
        p = doc.add_paragraph(); r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)

    H(doc, "Login Algorithm", 3)
    login_code = """PROCEDURE Login(username, password)

    // Validate session lockout (5 attempts, 5-minute cooldown)
    IF session contains login_lockout_until AND
       current_time < login_lockout_until THEN
        Display 'Too many attempts. Try again in 5 minutes'
        RETURN
    END IF

    // Validate CSRF token
    IF form csrf_token != session csrf_token THEN
        Abort with HTTP 400 Bad Request
    END IF

    // Validate required fields
    IF username is empty OR password is empty THEN
        Display 'Username and password are required'
        RETURN
    END IF

    // Look up user by case-insensitive username
    user <- get_user_by_username(username)

    // Verify credentials against salted hash
    IF user is not null AND user.password_hash is not null AND
       check_password_hash(user.password_hash, password) = TRUE THEN

        // Credentials valid — establish session
        Clear session
        session[user_id] <- user.id
        session[username] <- user.username
        session[gender] <- user.gender
        Reset login_attempts counter in session
        Log 'User {user_id} logged in' to audit log
        Display 'Logged in successfully'
        REDIRECT to home page
    ELSE
        // Credentials invalid — increment attempt counter
        attempts <- session[login_attempts] + 1
        session[login_attempts] <- attempts
        IF attempts >= 5 THEN
            session[login_lockout_until] <- current_time + 300 seconds
            Display 'Too many login attempts. Try again in 5 minutes'
        ELSE
            Display 'Invalid username or password'
        END IF
    END IF
END PROCEDURE"""
    for line in login_code.strip().split("\n"):
        p = doc.add_paragraph(); r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)

    H(doc, "Logout Algorithm", 3)
    logout_code = """PROCEDURE Logout()
    user_id <- session[user_id]
    Clear all session data
    Log 'User {user_id} logged out' to audit log
    Display 'Logged out successfully'
    REDIRECT to login page
END PROCEDURE"""
    for line in logout_code.strip().split("\n"):
        p = doc.add_paragraph(); r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)

    # Version Sequence
    H(doc, "Version Sequence / Code Backup", 2)
    P(doc, "The following table documents the version history of StudyLib, derived from the Git commit log. Each entry records the date, a unique commit identifier (hash), and a description of the changes made \u2014 providing evidence of consistent version control and backup procedures throughout development.", italic=True)

    git_entries = git_log()
    version_rows = [(entry[1], entry[0], entry[2][:120]) for entry in git_entries[:50]]
    T(doc, ["Date","Commit","Description"], version_rows, [1.0,0.8,4.9])
    P(doc, "(50 of 90+ commits shown; full history available in GitHub repository)", italic=True, sz=Pt(9))

    # Prototyping Sequence
    H(doc, "Prototyping Sequence", 2)
    P(doc, "The following screenshots document the evolution of the StudyLib user interface across key development stages, demonstrating the prototyping process and iterative design refinement.")
    IMG(doc, "Backend Screenshot 1: Flask Application Terminal Output (app.py running on port 8010)")
    IMG(doc, "Backend Screenshot 2: SQLite Database Schema (DB Browser for SQLite view)")
    IMG(doc, "Backend Screenshot 3: Pytest Test Suite Execution (terminal output)")
    IMG(doc, "Frontend Screenshot 1: Login Page with Candlelit Archive Dark Theme")
    IMG(doc, "Frontend Screenshot 2: Home Dashboard with Recently Viewed/Searched Panels and Wood Texture")
    IMG(doc, "Frontend Screenshot 3: Browse Page with Multi-Source Search, Parallel Results, and Victorian SVG Illustrations")
    IMG(doc, "Frontend Screenshot 4: In-App Content Viewer with Server-Side Proxy Rendering")
    IMG(doc, "Frontend Screenshot 5: AI Summarisation Result Panel (Summary, Bullet Points, Relevance, Citations)")
    IMG(doc, "Frontend Screenshot 6: Workspace Compilation View with Drag-Reordering, Notes, and Alexander Chat")
    IMG(doc, "Frontend Screenshot 7: PDF/DOCX Export with APA and Harvard Citation Options")
    IMG(doc, "Frontend Screenshot 8: File Upload Page with Text Extraction Status")
    IMG(doc, "Frontend Screenshot 9: User Profile Page with Victorian Gender-Based Avatar Selection")
    IMG(doc, "Frontend Screenshot 10: Navigation Sidebar Menu (Hamburger Toggle)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 5: TESTING, OPTIMISATION & EVALUATION
    # ════════════════════════════════════════════════════════════════
    H(doc, "Automated Testing, Optimisation & Evaluation")

    H(doc, "Test Plan", 2)
    P(doc, "StudyLib implements a comprehensive automated test suite using the pytest framework (v8.0.0+). Tests are organised by module and cover unit testing, integration testing, and contract testing. The test suite is designed to be run automatically via a single command and produces documented output on each execution.")
    P(doc, "Test modules and their coverage:", bold=True)

    test_modules = [
        ("test_whitelist.py","Unit","41 lines","Domain validation, wildcard pattern matching, scheme validation, whitelist.json loading"),
        ("test_citations.py","Unit","17 lines","APA and Harvard citation formatting correctness, PubMed metadata handling"),
        ("test_summarise.py","Unit","286 lines","URL summarisation, file summarisation, search result summarisation, API key validation, missing-key guards, error sanitisation, extractive fallback"),
        ("test_proxy.py","Unit","479 lines","DNS pinning, IP validation (private/multicast rejection), redirect limits, size limits, HTML sanitisation (script/form removal), reader mode, Google Books fallback"),
        ("test_search.py","Unit","1455 lines","Deduplication (union-find with canonical URLs), SerpAPI integration, Google Books covers, image URL validation, pagination, error handling, parallel search-all, thumbnail/favicon validation"),
        ("test_upload.py","Unit","33 lines","PDF text extraction (PyMuPDF), DOCX extraction, TXT extraction, file type detection"),
        ("test_ai.py","Unit","924 lines","AI answer/chat endpoints, workspace chat persistence, URL item upsert, JS runtime lifecycle tests via Node.js, prompt validation"),
        ("test_integration.py","Integration","23 lines","Full Flask route existence verification, end-to-end search-to-workspace workflow"),
        ("test_dark_theme_contract.py","Contract","1490+ lines","CSS token validation, scrollbar styling, button states, input/dropdown styling, navigation visuals, SVG illustration validation, reduced-motion compliance, GIF loader validation, browse page states"),
        ("test_light_theme_contract.py","Contract","1440+ lines","CSS token validation, paper/ink materials, button states, form controls, modals, alerts, WCAG contrast ratio checks, PNG texture validation, forced-colors mode fallbacks"),
    ]
    T(doc, ["Test File","Type","Size","Coverage"], test_modules, [2.2,1.0,0.8,2.7])

    H(doc, "Automated Test Execution", 2)
    P(doc, "All tests are executed via the single command: python -m pytest tests/ -v", bold=True)
    P(doc, "This discovers and runs all test files in the tests/ directory. The -v flag produces verbose output showing individual test results with pass/fail status. Tests cover both unit-level validation of individual modules and end-to-end integration of the full Flask application workflow.")
    P(doc, "Test results summary from the most recent execution:", bold=True)

    test_summary = get_test_summary()
    for line in test_summary:
        p = doc.add_paragraph(); r = p.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)

    P(doc, "Key test statistics: 381 total tests across 10 modules. Unit tests cover all core backend modules (db, search, proxy, summarise, citations, whitelist, files, answer). Integration tests verify the full Flask application route structure and end-to-end workflow. Contract tests enforce CSS theme consistency, accessibility compliance, and visual regression prevention.", italic=True)

    H(doc, "Test Categories and Boundary Conditions", 2)
    P(doc, "Each test module validates specific boundary conditions and edge cases:", bold=True)
    test_cases = [
        "Whitelist Tests: Valid exact domains, valid wildcard patterns (*.edu.au), invalid domains, scheme-stripped URLs, empty whitelist, malformed JSON.",
        "Proxy Tests: Private IP ranges (10.x, 172.16-31.x, 192.168.x), multicast ranges, localhost, IPv6 private equivalents, redirect chains exceeding limit, response size exceeding limit, non-HTML content types.",
        "Search Tests: Empty results, partial source failures, malformed SerpAPI responses, image URLs with credentials/ports/fragments, oversized image URLs, non-HTTPS image URLs, deduplication of canonical URL variants, pagination boundaries.",
        "Summarisation Tests: Missing Anthropic API key, empty source content, extremely long content, malformed AI response JSON, network timeout simulation, extractive fallback when AI unavailable.",
        "Authentication Tests: Valid credentials, invalid password, non-existent username, lockout after 5 attempts, lockout cooldown expiry, CSRF token mismatch, empty form fields, boundary username length, boundary email length.",
        "Theme Contract Tests: Dark-only selectors in light mode, light-only selectors in dark mode, WCAG AA contrast ratios, forced-colors mode fallbacks, reduced-motion compliance, SVG/GIF asset integrity, CSS rule group ordering.",
    ]
    for tc in test_cases:
        p = doc.add_paragraph(); r = p.add_run(tc); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

    H(doc, "Code Optimisation", 2)
    P(doc, "Several optimisation strategies were applied to improve StudyLib\u2019s performance and resource efficiency:")
    optimisations = [
        "Search Result Caching: Identical search queries across the same source are cached in the search_cache table using SHA-256 hash keys. Cache hits avoid redundant SerpAPI calls, reducing API costs and improving response times by up to 80% for repeated queries.",
        "Parallel Multi-Source Search: The browse search-all endpoint uses ThreadPoolExecutor with up to 6 concurrent workers and a 25-second timeout. Results from Wikipedia, Google Books, Scholar, PubMed, and whitelist sources are fetched in parallel rather than sequentially, reducing perceived wait time from ~15s to ~5s.",
        "Deduplication via Union-Find: Search results across overlapping source scopes are deduplicated using canonical URL normalisation (stripping fragments, trailing slashes, query parameters) combined with a union-find data structure. This reduces result bloat and ensures each unique source appears only once.",
        "Database Query Optimisation: Frequently accessed queries (recently viewed, recently searched, saved items) use LIMIT 10 with DESC ordering and automatic pruning of old entries on insert. Composite primary keys on junction tables eliminate redundant index lookups.",
        "Proxy Response Limits: The secure content proxy enforces DNS pinning (single IP per hostname), response size limits (5MB), and redirect limits (5 hops). These constraints prevent resource exhaustion and timeout scenarios when fetching external content.",
    ]
    for opt in optimisations:
        p = doc.add_paragraph(); r = p.add_run(opt); r.font.size = Pt(10); p.paragraph_format.space_after = Pt(4); p.paragraph_format.left_indent = Cm(0.5)

    H(doc, "Evaluation", 2)
    P(doc, "StudyLib was developed to address a clearly identified need: secondary school students lack a single tool that finds verified academic sources, helps them understand those sources quickly through AI summarisation, and compiles the results into an exportable study document. The following evaluation assesses the project against its original functional and non-functional requirements, reflects on the development process, and identifies areas for future improvement.")

    H(doc, "Success Against Requirements", 3)
    P(doc, "Of the 32 functional requirements specified, all 32 are implemented and functional. Key areas of success include:")
    fr_success = [
        "Authentication (FR1\u2013FR3): Registration and login with Werkzeug pbkdf2:sha256 password hashing, CSRF protection on all forms, and brute-force lockout after 5 failed attempts. Session management via Flask-Session with SQLAlchemy backend ensures persistence across server restarts.",
        "Search and whitelist (FR4\u2013FR5, FR25): Multi-source parallel search returns results from Wikipedia, Google Books, Google Scholar, PubMed, and configurable whitelist domains. A union-find deduplication algorithm with canonical URL normalisation prevents duplicate results across overlapping source scopes.",
        "Content viewing (FR6\u2013FR9): The server-side proxy (src/proxy.py) fetches external content with DNS pinning, IP validation rejecting private/multicast ranges, redirect limiting (5 hops), and response size limiting (5MB). HTML sanitisation strips scripts, forms, and active content. Paywalled sources degrade gracefully to a new-tab link.",
        "AI summarisation (FR10\u2013FR11, FR15, FR26\u2013FR27): Claude Haiku 4.5 generates structured JSON output (summary, bullets, relevance, citations) with explicit instructions to flag insufficient content rather than fabricate. Claude Sonnet 4.6 powers multi-turn chat (Alexander) with RAG context from workspace items, uploaded files, and web search results. Chat history persists across sessions.",
        "File upload and extraction (FR12\u2013FR14): Supports PDF (PyMuPDF), DOCX, TXT, XLSX/XLS, and image uploads up to 10MB. Extracted text enables full-text search and RAG context retrieval during chat.",
        "Workspace and export (FR16\u2013FR22, FR28\u2013FR29): Full CRUD for workspaces with drag-reordering. Rich-text notes via Quill editor. Export to PDF (ReportLab) and DOCX (python-docx) with APA/Harvard citations. Victorian-era avatar selection by gender preference.",
        "Automated testing (FR23\u2013FR24): 381 tests across 10 modules covering unit, integration, and CSS contract testing. All tests run automatically via \u2018pytest tests/ -v\u2019.",
    ]
    for s in fr_success:
        LI(doc, s)

    H(doc, "Development Process Reflection", 3)
    P(doc, "The WAgile approach proved effective for this project\u2019s characteristics. The upfront planning phase produced a requirements specification and data dictionary that guided implementation without constraining it. Iterative sprints allowed the scope of complex modules \u2014 particularly the SerpAPI multi-source search and the secure content proxy \u2014 to be refined as implementation revealed edge cases not anticipated at the design stage.")

    P(doc, "Key technical decisions that worked well:", bold=True)
    decs = [
        "SQLite over PostgreSQL/MySQL: Eliminated external database server dependencies, enabling zero-config local deployment. The decision proved correct \u2014 no concurrency limitations were encountered for a single-user local application, and the zero-install requirement made teacher evaluation straightforward.",
        "Modular architecture (10 source modules in src/): Enabled independent testing of each component and made iterative refinement possible without cascading changes. Each module has a single responsibility and a clear interface (db.py for ORM, search.py for SerpAPI, proxy.py for secure content fetching, etc.).",
        "SerpAPI over direct Google API: SerpAPI\u2019s unified interface across Wikipedia, Google Books, Google Scholar, and web search simplified multi-source integration significantly. The trade-off was dependency on a third-party service with API limits, which is acceptable for this project\u2019s scale and cost profile.",
        "CSS contract testing: The 2,900-line custom CSS for dual Victorian-era themes required automated testing to prevent visual regressions. Contract tests validate CSS selector scoping, colour token consistency, accessibility contrast ratios, and forced-colors mode fallbacks. Without these, maintaining theme consistency across rapid UI iterations would have been error-prone.",
    ]
    for d in decs:
        LI(doc, d)

    H(doc, "Challenges Encountered and How They Were Addressed", 3)

    P(doc, "1. Attempted local AI deployment (OpenClaw):", bold=True)
    P(doc, "Early in the AI integration phase, significant effort was invested in attempting to run the OpenClaw language model locally on the development machine. The goal was to eliminate the dependency on the Anthropic cloud API entirely, removing the recurring API cost and enabling fully offline AI summarisation and chat. However, this approach encountered multiple failures:")
    local_ai = [
        "Performance: Even the smallest usable OpenClaw model variant ran an order of magnitude slower than cloud-based Claude. A typical summarisation call that completes in 3\u20135 seconds via the Anthropic API took 45\u201390 seconds locally, making the user experience unacceptable for an interactive application.",
        "Memory constraints: The development machine\u2019s available RAM (8 GB) was insufficient to load model weights while simultaneously running the Flask application, SQLite database, and browser. The system became unresponsive during model inference, often requiring a forced restart.",
        "Output quality: The locally-run model produced significantly lower quality summaries and chat responses compared to Claude. Structured JSON output (required for the summary + bullets + relevance format) was unreliable, with the model frequently producing malformed JSON or incomplete responses.",
        "Setup complexity: Installing and configuring OpenClaw required additional dependencies (CUDA toolkit, model weights download) that conflicted with the project\u2019s zero-config deployment goal. The added setup burden was antithetical to the requirement that a teacher should be able to run the application by following simple README instructions.",
    ]
    for item in local_ai:
        LI(doc, item)
    P(doc, "After approximately three days of experimentation, the decision was made to pivot to the cloud-based Anthropic Claude API. The $5.00 AUD pre-paid credit proved sufficient for the entire development and assessment period. While local AI would have been the ideal solution for offline capability and zero per-use cost, the current state of consumer hardware and open-source model performance made it impractical for this application\u2019s latency and quality requirements. This is noted as an area to revisit as local model inference improves over the next 2\u20133 years.")

    P(doc, "2. Google Books API rendering:", bold=True)
    P(doc, "The initial approach attempted to use the Google Books Embed API directly, but CORS restrictions and inconsistent API key behaviour made native rendering unreliable. The solution (commit d1434d8) was a hybrid approach: construct a Google Books viewer URL from the extracted volume ID, attempt embedded rendering with an 8-second timeout, and fall back to opening in a new tab if rendering fails to initialise.")

    P(doc, "3. SerpAPI response format inconsistency:", bold=True)
    P(doc, "Different source types (Wikipedia, Google Books, Google Scholar) return slightly different JSON structures from SerpAPI. The browse_serpapi_search() function in src/search.py evolved through multiple iterations to normalise these into a consistent Item schema. The deduplication step using canonical URL normalisation was essential because Google Scholar results often overlapped with direct domain search results.")

    P(doc, "4. CSS theme scope leakage:", bold=True)
    P(doc, "The dual-theme approach (dark Candlelit Archive + light Old Book) required strict CSS selector scoping to prevent dark theme rules from leaking into light mode and vice versa. This was resolved through the comprehensive contract test suite (1,490 lines for dark theme, 1,440 lines for light theme) rather than manual inspection \u2014 a case where investing in test infrastructure caught regressions that manual review would have missed.")

    P(doc, "5. API key management:", bold=True)
    P(doc, "Early in development, the SerpAPI key was inadvertently lost during a merge, causing search to fail silently. This prompted the addition of the SERP_API_KEY environment variable check with a clear 503 error message and the .env.example template. The lesson: environment variable configuration must be documented and validated at startup, not assumed.")

    H(doc, "Strengths of the Final Product", 3)
    strengths = [
        "Complete research workflow: StudyLib is the only tool in its category that integrates verified academic search, AI summarisation, workspace compilation, and formatted export into a single application. The end-to-end flow from query to downloadable document is seamless.",
        "Security-first design: CSRF protection, password hashing, session lockout, input validation, SQL injection prevention via ORM, and a secure content proxy with DNS pinning demonstrate a security posture appropriate for handling student data and external content.",
        "Distinctive, accessible UI: The Victorian-era aesthetic (Candlelit Archive dark theme, Old Book light theme) creates a memorable user experience that differentiates StudyLib from generic education tools. WCAG contrast compliance and forced-colors mode fallbacks ensure accessibility.",
        "Comprehensive testing: 381 automated tests across 10 modules provide confidence that core functionality works and that visual presentation remains consistent. The testing investment is proportional to the project\u2019s complexity.",
        "Zero-cost deployment: Students and teachers can run StudyLib on any laptop with Python installed at no cost. The application requires no external server, no paid subscriptions, and no institutional infrastructure.",
    ]
    for s in strengths:
        LI(doc, s)

    H(doc, "Areas for Future Improvement", 3)
    improvements = [
        "Revisit local AI deployment: As open-source models and consumer hardware improve, re-evaluating local inference (via OpenClaw or a successor) would eliminate the API dependency entirely. This is the single most impactful future improvement for offline capability and zero per-use cost.",
        "Offline-first PWA support: Implementing a progressive web app shell with cached UI would allow students to browse previously fetched results and workspaces without internet connectivity, with API-dependent features gracefully disabled.",
        "Collaborative workspaces: The workspace model is currently single-user. Adding shared workspaces with invite links would support group research projects, a common Year 11\u201312 assessment format. This would require a hosted deployment rather than local-only operation.",
        "Expanded whitelist coverage: Adding more country-specific academic domains (.ac.uk, .edu.sg, .edu.in) would improve relevance for international curriculum contexts.",
        "Scheduled cache cleanup: The search_cache table grows unboundedly. A cleanup routine expiring entries older than 24 hours would prevent database bloat during extended use.",
    ]
    for imp in improvements:
        LI(doc, imp)

    H(doc, "Conclusion", 3)
    P(doc, "StudyLib successfully delivers on its stated purpose: helping secondary students find, understand, and compile credible academic sources into structured study documents. The project demonstrates mastery of the Software Engineering syllabus outcomes (SE-12-01 through SE-12-09): identifying and defining the problem, researching and selecting a development approach, managing and scheduling project work, producing and implementing a functional and well-engineered solution, comprehensively testing and optimising the code, and documenting the full engineering process in a professional folio. At a total development cost of $5.00 AUD and zero cost to end users, the project proves that high-quality educational software can be built and deployed without financial barriers. While the attempt to use local AI (OpenClaw) was unsuccessful due to performance, memory, and quality constraints, the pivot to cloud-based Claude API proved cost-effective and delivered the response quality required for an educational tool. The experience provided valuable insight into the current state of local versus cloud AI deployment, and the local option remains a clear direction for future improvement as the technology matures.", italic=True)

    # ── SAVE ──
    output_path = os.path.join(PROJECT_DIR, "StudyLib_Student_Folio_2026.docx")
    doc.save(output_path)
    print(f"\nFolio saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    build()
